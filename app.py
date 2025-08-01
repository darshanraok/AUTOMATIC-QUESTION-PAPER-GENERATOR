from docx import Document
from docx.shared import Pt
import sqlite3
import random
from datetime import datetime
from docx2pdf import convert
import os

def replace_top_text(doc, subject):
    """Replace the top text in the document based on the selected subject."""
    for para in doc.paragraphs:
        if "22MCA101" in para.text or "22MCA204" in para.text or "22MCA301" in para.text or "MMCA21" in para.text or "MMCA23" in para.text or "MMCA24" in para.text or "MMCA25" in para.text or "MMCA261" in para.text or "MMMCA262" in para.text:
            para.clear()
            title_run = set_font(para, 'Times New Roman', 13, bold=True)
            if subject == "java":
                title_run.text = "Course Code: 22MCA204"
            elif subject == "machine learning":
                title_run.text = "Course Code: 22MCA301"
            elif subject == "full stack development":
                title_run.text = "Course Code: MMCA21"
            elif subject == "mobile application development":
                title_run.text = "Course Code: MMCA23"
            elif subject == "cloud computing":
                title_run.text = "Course Code: MMCA24"
            elif subject == "computer networks":
                title_run.text = "Course Code: MMCA25"
            elif subject == "cyber security":
                title_run.text = "Course Code: MMCA261"
            elif subject == "robotic process automation(rpa)":
                title_run.text = "Course Code: MMCA262"
            else: # mfca
                title_run.text = "Course Code: 22MCA101"
            break

def extract_10_mark_questions(db_path, module_name):
    """Function to extract 10-mark questions from a specific module in the database."""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    query = f"SELECT question_name, marks FROM {module_name} WHERE marks = 10"
    cursor.execute(query)
    questions = cursor.fetchall()
    conn.close()
    return questions

def set_font(para, font_name, size, bold=False):
    """Function to set font style."""
    run = para.add_run()
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    return run

def replace_keywords_in_doc(doc, questions_by_module):
    """Function to replace keywords in the DOCX with questions and set font, ensuring no repeats."""
    # Create a copy of the question lists to modify in memory
    questions_to_use = {key: list(value) for key, value in questions_by_module.items()}

    for para in doc.paragraphs:
        for keyword, questions in questions_to_use.items():
            if questions and keyword in para.text:
                question = random.choice(questions)
                question_text = question[0]
                para.clear()
                run = para.add_run(question_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                questions.remove(question) # Remove the used question from the list

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for keyword, questions in questions_to_use.items():
                    if questions and keyword in cell.text:
                        question = random.choice(questions)
                        question_text = question[0]
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(question_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        questions.remove(question) # Remove the used question from the list


def replace_co_text(doc, subject):
    """Function to replace CO text based on subject and set font/alignment."""
    co_replacements = {}
    if subject == 'mfca':
        co_replacements = {
            "C011": "Apply the concepts of set theory and propositional logic to solve problems.",
            "C022": "Compute statistical measures of random variables and probability distributions.",
            "C033": "Solve problems using concepts of relations.",
            "C044": "Apply the abstract concepts of graph theory to solve problems.",
            "C055": "Fit an appropriate curve for the given data."
        }
    elif subject == 'java':
        co_replacements = {
            "C011": "Demonstrate the basic programming constructs of Java and OOP concepts to develop Java applications.",
            "C022": "Illustrate the concepts of generalization and run time polymorphism to develop reusable components.",
            "C033": "Exemplify the usage of Multithreading in building efficient applications.",
            "C044": "Build web applications using Servlets and JSP.",
            "C055": "Design applications using JDBC and Enterprise Java Beans."
        }
    elif subject == 'machine learning':
        co_replacements = {
            "C011": "Explore the Machine Learning concepts.",
            "C022": "Build suitable Decision tree for a given data set.",
            "C033": "Apply machine learning algorithms for the given problems.",
            "C044": "Perform statistical and probabilistic analysis of machine learning techniques.",
            "C055": "Implement machine learning algorithms for a given use case."
        }
    elif subject == 'full stack development':
        co_replacements = {
            "C011": "Demonstrate MERN stack and its components.",
            "C022": "Design React applications using various components.",
            "C033": "Build dynamic web applications.",
            "C044": "Implement RESTful APIs.",
            "C055": "Integrate MongoDB for data management."
        }
    elif subject == 'mobile application development':
        co_replacements = {
            "C011": "Analyze the features of mobile devices.",
            "C022": "Design applications using Android components.",
            "C033": "Develop mobile based application using database.",
            "C044": "Design an application using image capturing and location details.",
            "C055": "Create a mobile application for any given use case."
        }
    elif subject == 'cloud computing':
        co_replacements = {
            "C011": "Analyse the requirements for scalable services and computing environment.",
            "C022": "Classify various cloud service models and their providers.",
            "C033": "Compare various cloud deployment models.",
            "C044": "Deploy applications on real-time cloud platforms.",
            "C055": "Get any 2 cloud certificates (aws+azure) or (aws+googlecloud)."
        }
    elif subject == 'computer networks':
        co_replacements = {
            "C011": "Apply the basic concepts of computer networking.",
            "C022": "Demonstrate OSI reference model and TCP/IP model.",
            "C033": "Analyze the working of network protocols.",
            "C044": "Implement networking concepts using appropriate tools.",
            "C055": "Become Network engineer or Network architect."
        }
    elif subject == 'cyber security':
        co_replacements = {
            "C011": "Analyse cyber security threats and vulnerabilities.",
            "C022": "Apply cryptographic techniques to secure the data.",
            "C033": "Demonstrate system resilience to cyber-attacks.",
            "C044": "Develop incident response plans with forensic analyses",
            "C055": "Always be carefull when sharing your data."
        }
    elif subject == 'robotic process automation(rpa)':
        co_replacements = {
            "C011": "Analyse the problem to understand the scope and extent of process automation.",
            "C022": "Apply the robotic process automation knowledge to automate operations.",
            "C033": "Implement exception handling and automation strategies in real time applications.",
            "C044": "Interpret various aspects of debugging in RPA applications.",
            "C055": "Develop basic robots using UiPath Community Edition."
        }
    else:
        return

    for para in doc.paragraphs:
        for co, replacement_text in co_replacements.items():
            if co in para.text:
                para.clear()
                run = para.add_run(replacement_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for co, replacement_text in co_replacements.items():
                    if co in cell.text:
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(replacement_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)


def main():
    try:
        print("Choose the subject:")
        print("1. MFCA (22MCA101)")
        print("2. Java (22MCA204)")
        print("3. Machine Learning (22MCA301)")
        print("4. Full Stack Development (MMCA21)")
        print("5. Mobile Application Development (MMCA23)")
        print("6. Cloud Computing (MMCA24)")
        print("7. Computer Networks (MMCA25)")
        print("8. Cyber Security (MMCA261)")
        print("9. Robotic Process Automation (MMCA262)")
        print("10. Exit")

        choice = input("Enter your choice (1-10): ").strip()

        subject_map = {
            "1": "mfca",
            "2": "java",
            "3": "machine learning",
            "4": "full stack development",
            "5": "mobile application development",
            "6": "cloud computing",
            "7": "computer networks",
            "8": "cyber security",
            "9": "robotic process automation(rpa)"
        }

        subject = subject_map.get(choice)

        if subject is None:
            if choice == "10":
                print("Exiting the program.")
            else:
                print("Invalid choice. Please enter a number between 1 and 10.")
            return

        db_path = f'./database/{subject.replace(" ", "").replace("(","").replace(")","")}.db'

        if not os.path.exists(db_path):
            print(f"Database file not found for {subject}. Please ensure '{subject.replace(' ', '').replace('(','').replace(')','')}.db' exists in the 'database' directory.")
            return

        # Fetch all questions once and store them in a dictionary
        questions_by_module = {
            "M11": extract_10_mark_questions(db_path, 'module1'), "M12": extract_10_mark_questions(db_path, 'module1'),
            "M13": extract_10_mark_questions(db_path, 'module1'), "M14": extract_10_mark_questions(db_path, 'module1'),
            "M15": extract_10_mark_questions(db_path, 'module1'),
            "M21": extract_10_mark_questions(db_path, 'module2'), "M22": extract_10_mark_questions(db_path, 'module2'),
            "M23": extract_10_mark_questions(db_path, 'module2'), "M24": extract_10_mark_questions(db_path, 'module2'),
            "M31": extract_10_mark_questions(db_path, 'module3'), "M32": extract_10_mark_questions(db_path, 'module3'),
            "M33": extract_10_mark_questions(db_path, 'module3'), "M34": extract_10_mark_questions(db_path, 'module3'),
            "M41": extract_10_mark_questions(db_path, 'module4'), "M42": extract_10_mark_questions(db_path, 'module4'),
            "M43": extract_10_mark_questions(db_path, 'module4'), "M44": extract_10_mark_questions(db_path, 'module4'),
            "M51": extract_10_mark_questions(db_path, 'module5'), "M52": extract_10_mark_questions(db_path, 'module5'),
            "M53": extract_10_mark_questions(db_path, 'module5'), "M54": extract_10_mark_questions(db_path, 'module5'),
        }

        template_path = './pattern/pattern.docx'
        doc = Document(template_path)

        for para in doc.paragraphs:
            if "MATHEMATICAL FOUNDATION FOR COMPUTER APPLICATIONS" in para.text or "JAVA PROGRAMMING" in para.text or "MACHINE LEARNING" in para.text or "FULL STACK DEVELOPMENT" in para.text or "MOBILE APPLICATION DEVELOPMENT" in para.text or "CLOUD COMPUTING" in para.text or "COMPUTER NETWORKS" in para.text or "CYBER SECURITY" in para.text or "ROBOTIC PROCESS AUTOMATION(RPA)" in para.text:
                para.clear()
                title_run = set_font(para, 'Times New Roman', 16, bold=True)
                if subject == "java":
                    title_run.text = "JAVA PROGRAMMING"
                elif subject == "machine learning":
                    title_run.text = "MACHINE LEARNING"
                elif subject == "full stack development":
                    title_run.text = "FULL STACK DEVELOPMENT"
                elif subject == "mobile application development":
                    title_run.text = "MOBILE APPLICATION DEVELOPMENT"
                elif subject == "cloud computing":
                    title_run.text = "CLOUD COMPUTING"
                elif subject == "computer networks":
                    title_run.text = "COMPUTER NETWORKS"
                elif subject == "cyber security":
                    title_run.text = "CYBER SECURITY"
                elif subject == "robotic process automation(rpa)":
                    title_run.text = "ROBOTIC PROCESS AUTOMATION"
                else:
                    title_run.text = "MATHEMATICAL FOUNDATION FOR COMPUTER APPLICATIONS"
                break

        replace_keywords_in_doc(doc, questions_by_module)
        replace_co_text(doc, subject)
        replace_top_text(doc, subject)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_output_path = f'./generatedPapers/{subject.replace(" ", "_").replace("(","").replace(")","")}_{timestamp}.docx'

        doc.save(docx_output_path)

        pdf_output_path = f'./generatedPapers/{subject.replace(" ", "_").replace("(","").replace(")","")}_{timestamp}.pdf'
        convert(docx_output_path, pdf_output_path)

        print(f"Generated question paper saved as: {docx_output_path}")
        print(f"Generated question paper PDF saved as: {pdf_output_path}")

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
